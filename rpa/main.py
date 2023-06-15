# 날짜 : 2023년 06월 11일
# 프로젝트 : 성공회 성가 가사 수집 프로젝트
# 학습목표 : 파이썬, 웹 크롤링, RPA

# 경고메세지 끄기
import warnings
warnings.filterwarnings(action='ignore')
# 다시 출력하게 하기
#warnings.filterwarnings(action='default')

# HTML 파일에서 링크 수집
from lyrics_crawling import *
links, titles = ft_get_links("source.html")

## 1번부터 100번 링크 수집 -> 이미지 수집용
#index = 0
#link_number = len(links)
#while index < 100: # 원하는 개수 입력
#    index+=1
#    print(links[link_number - index])
    
# 링크 들어가서 가사 수집
lyrics = ft_get_lyrics(links)

# wd 파일로 출력
from docx import Document
doc = Document()
index = 0
len_link = len(links)
while index < len_link:
    index += 1
    doc.add_heading(titles[len_link - index])
    for lyric in lyrics[len_link - index]:
        doc.add_paragraph(lyric)
    doc.add_paragraph("")
doc.save('성공회 성가.docx')
print("파일 생성 완료")

# csv 파일로 출력
#import csv
#index = 0
#filename = "test.csv"
   
#f = open(filename, "w", encoding="utf8", newline="")
#writer = csv.writer(f)

#for lyric in lyrics:
#    writer.writerow(titles[index])
#    index +=1